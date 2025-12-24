"""
Intellibooks Studio - RAG Pipeline with Ray, RabbitMQ, and ChromaDB

This module implements a production-ready RAG pipeline with:
- Ray for distributed parallel processing of documents
- RabbitMQ for task queuing and async processing
- ChromaDB (Docker) for vector storage
- Sentence Transformers for embeddings
- LangChain for LLM integration
"""

import os
import io
import json
import hashlib
import logging
import asyncio
import time
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from uuid import uuid4
from concurrent.futures import ThreadPoolExecutor
from functools import lru_cache

import numpy as np
from pydantic import BaseModel

logger = logging.getLogger(__name__)

# ============ Configuration ============

@dataclass
class RAGConfig:
    """Configuration for RAG pipeline."""
    # ChromaDB settings
    chroma_host: str = "localhost"
    chroma_port: int = 8000
    chroma_collection: str = "intellibooks_knowledge"
    chroma_use_http: bool = True  # Use HTTP client (Docker) or persistent client (local)
    chroma_persist_directory: str = None  # Directory for persistent client (auto-set if None)

    # Embedding settings
    embedding_model: str = "all-MiniLM-L6-v2"
    embedding_dimension: int = 384

    # Chunking settings
    chunk_size: int = 1000
    chunk_overlap: int = 200

    # Processing settings
    max_workers: int = 4
    batch_size: int = 50

    # RabbitMQ settings (optional)
    rabbitmq_host: str = "localhost"
    rabbitmq_port: int = 5672
    rabbitmq_user: str = "admin"
    rabbitmq_password: str = "devpassword123"
    rabbitmq_queue: str = "rag_documents"
    use_rabbitmq: bool = False

    # Ray settings (optional) - connects to Docker Ray cluster
    use_ray: bool = False
    ray_address: str = "ray://localhost:10001"  # Ray client port from Docker
    ray_num_cpus: int = 4

    # LLM settings
    openrouter_api_key: str = ""
    openrouter_model: str = "anthropic/claude-sonnet-4"


def load_config() -> RAGConfig:
    """Load configuration from environment."""
    from dotenv import load_dotenv

    env_path = Path(__file__).parent.parent.parent.parent / ".env"
    if env_path.exists():
        load_dotenv(env_path)

    # Determine persist directory
    persist_dir = os.getenv("CHROMA_PERSIST_DIRECTORY")
    if not persist_dir:
        persist_dir = str(Path(__file__).parent.parent.parent.parent / "data" / "chroma_db")
    
    return RAGConfig(
        chroma_host=os.getenv("CHROMA_HOST", "localhost"),
        chroma_port=int(os.getenv("CHROMA_PORT", "8000")),
        chroma_collection=os.getenv("CHROMA_COLLECTION", "intellibooks_knowledge"),
        chroma_use_http=os.getenv("CHROMA_USE_HTTP", "true").lower() == "true",
        chroma_persist_directory=persist_dir,
        embedding_model=os.getenv("EMBEDDING_MODEL", "all-MiniLM-L6-v2"),
        chunk_size=int(os.getenv("RAG_CHUNK_SIZE", "1000")),
        chunk_overlap=int(os.getenv("RAG_CHUNK_OVERLAP", "200")),
        max_workers=int(os.getenv("RAG_MAX_WORKERS", "4")),
        rabbitmq_host=os.getenv("RABBITMQ_HOST", "localhost"),
        rabbitmq_port=int(os.getenv("RABBITMQ_PORT", "5672")),
        rabbitmq_user=os.getenv("RABBITMQ_USER", "admin"),
        rabbitmq_password=os.getenv("RABBITMQ_PASSWORD", "devpassword123"),
        use_rabbitmq=os.getenv("USE_RABBITMQ", "false").lower() == "true",
        use_ray=os.getenv("USE_RAY", "false").lower() == "true",
        ray_address=os.getenv("RAY_ADDRESS", "ray://localhost:10001"),
        openrouter_api_key=os.getenv("OPENROUTER_API_KEY", ""),
        openrouter_model=os.getenv("OPENROUTER_MODEL", "anthropic/claude-sonnet-4"),
    )


# ============ Data Models ============

class Document(BaseModel):
    """A document chunk for indexing."""
    id: str
    content: str
    metadata: Dict[str, Any] = {}


class SearchResult(BaseModel):
    """A search result from vector store."""
    id: str
    content: str
    score: float
    metadata: Dict[str, Any] = {}


@dataclass
class ProcessingResult:
    """Result of document processing."""
    success: bool
    document_id: str
    filename: str
    chunks_created: int
    processing_time_ms: float
    error: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class RAGResponse:
    """Response from RAG query."""
    answer: str
    sources: List[Dict[str, Any]]
    query: str
    processing_time_ms: float
    confidence: float = 0.0


# ============ Embedding Service ============

class EmbeddingService:
    """Service for generating embeddings."""

    _instance = None
    _model = None

    def __new__(cls, model_name: str = "all-MiniLM-L6-v2"):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance._model_name = model_name
        return cls._instance

    @property
    def model(self):
        """Lazy load embedding model."""
        if self._model is None:
            try:
                from sentence_transformers import SentenceTransformer
                self._model = SentenceTransformer(self._model_name)
                logger.info(f"Loaded embedding model: {self._model_name}")
            except ImportError:
                logger.warning("SentenceTransformer not available, using fallback")
                self._model = FallbackEmbedder()
        return self._model

    def encode(self, texts: List[str]) -> np.ndarray:
        """Generate embeddings for texts."""
        return self.model.encode(texts, convert_to_numpy=True)

    def encode_single(self, text: str) -> np.ndarray:
        """Generate embedding for single text."""
        return self.encode([text])[0]


class FallbackEmbedder:
    """Fallback embedder when SentenceTransformer is unavailable."""

    def __init__(self, dim: int = 384):
        self.dim = dim

    def encode(self, texts: List[str], convert_to_numpy: bool = True) -> np.ndarray:
        embeddings = []
        for text in texts:
            np.random.seed(hash(text) % (2**32))
            emb = np.random.randn(self.dim).astype(np.float32)
            emb = emb / np.linalg.norm(emb)
            embeddings.append(emb)
        return np.array(embeddings)


# ============ ChromaDB Vector Store (Docker) ============

class ChromaDBHttpClient:
    """Direct HTTP client for ChromaDB Docker instance.

    This bypasses the chromadb Python package which has Pydantic v1/v2 compatibility issues.
    Uses httpx to communicate directly with ChromaDB's REST API v2.
    """

    def __init__(self, host: str = "localhost", port: int = 8000):
        self.base_url = f"http://{host}:{port}"
        self._tenant = "default_tenant"
        self._database = "default_database"
        self._collection_ids = {}  # Cache collection name -> id mapping

    def _api_base(self) -> str:
        """Get the base path for tenant/database scoped API calls."""
        return f"/api/v2/tenants/{self._tenant}/databases/{self._database}"

    def _request(self, method: str, path: str, **kwargs) -> dict:
        """Make HTTP request to ChromaDB."""
        import httpx
        url = f"{self.base_url}{path}"
        with httpx.Client(timeout=60.0) as client:
            response = client.request(method, url, **kwargs)
            response.raise_for_status()
            return response.json() if response.content else {}

    def heartbeat(self) -> int:
        """Check if ChromaDB is alive."""
        result = self._request("GET", "/api/v2/heartbeat")
        return result.get("nanosecond heartbeat", 0)

    def get_or_create_collection(self, name: str, metadata: dict = None) -> str:
        """Get or create a collection, returns collection ID."""
        # Check cache first
        if name in self._collection_ids:
            return self._collection_ids[name]

        # Try to get existing collection first
        try:
            result = self._request(
                "GET",
                f"{self._api_base()}/collections/{name}"
            )
            collection_id = result.get("id", name)
            self._collection_ids[name] = collection_id
            return collection_id
        except Exception:
            pass

        # Create new collection
        result = self._request(
            "POST",
            f"{self._api_base()}/collections",
            json={
                "name": name,
                "metadata": metadata or {},
                "get_or_create": True
            }
        )
        collection_id = result.get("id", name)
        self._collection_ids[name] = collection_id
        return collection_id

    def _get_collection_id(self, collection_name: str) -> str:
        """Get collection ID from name, using cache."""
        if collection_name not in self._collection_ids:
            self.get_or_create_collection(collection_name)
        return self._collection_ids.get(collection_name, collection_name)

    def upsert(self, collection_name: str, ids: list, embeddings: list,
               documents: list, metadatas: list):
        """Upsert documents to collection."""
        collection_id = self._get_collection_id(collection_name)
        self._request(
            "POST",
            f"{self._api_base()}/collections/{collection_id}/upsert",
            json={
                "ids": ids,
                "embeddings": embeddings,
                "documents": documents,
                "metadatas": metadatas,
            }
        )

    def query(self, collection_name: str, query_embeddings: list, n_results: int = 5,
              where: dict = None, include: list = None) -> dict:
        """Query collection for similar documents."""
        collection_id = self._get_collection_id(collection_name)
        body = {
            "query_embeddings": query_embeddings,
            "n_results": n_results,
            "include": include or ["documents", "metadatas", "distances"],
        }
        if where:
            body["where"] = where

        return self._request(
            "POST",
            f"{self._api_base()}/collections/{collection_id}/query",
            json=body
        )

    def delete(self, collection_name: str, where: dict):
        """Delete documents matching filter."""
        collection_id = self._get_collection_id(collection_name)
        self._request(
            "POST",
            f"{self._api_base()}/collections/{collection_id}/delete",
            json={"where": where}
        )

    def delete_collection(self, collection_name: str):
        """Delete entire collection."""
        # ChromaDB v2 API requires collection NAME for deletion, not UUID
        self._request(
            "DELETE",
            f"{self._api_base()}/collections/{collection_name}"
        )
        # Clear from cache
        self._collection_ids.pop(collection_name, None)

    def count(self, collection_name: str) -> int:
        """Count documents in collection."""
        collection_id = self._get_collection_id(collection_name)
        result = self._request(
            "GET",
            f"{self._api_base()}/collections/{collection_id}/count"
        )
        return result if isinstance(result, int) else 0


class ChromaDBPersistentClient:
    """ChromaDB persistent client wrapper for local Python package."""
    
    def __init__(self, persist_directory: str, collection_name: str):
        try:
            import chromadb
            self.chromadb = chromadb
        except ImportError:
            raise ImportError("ChromaDB package not installed. Install with: pip install chromadb")
        
        self.persist_directory = persist_directory
        self.collection_name = collection_name
        self._client = None
        self._collection = None
        self._collection_ids = {collection_name: collection_name}  # For compatibility
        
        # Ensure directory exists
        Path(persist_directory).mkdir(parents=True, exist_ok=True)
        
    @property
    def client(self):
        """Get or create persistent client."""
        if self._client is None:
            self._client = self.chromadb.PersistentClient(path=self.persist_directory)
        return self._client
    
    @property
    def collection(self):
        """Get or create collection."""
        if self._collection is None:
            self._collection = self.client.get_or_create_collection(
                name=self.collection_name,
                metadata={"hnsw:space": "cosine"},
            )
        return self._collection
    
    def heartbeat(self) -> int:
        """Heartbeat check (always succeeds for persistent client)."""
        return 1
    
    def get_or_create_collection(self, name: str, metadata: dict = None) -> str:
        """Get or create collection (for compatibility with HTTP client interface)."""
        if name != self.collection_name:
            # Create a new collection if different name requested
            collection = self.client.get_or_create_collection(
                name=name,
                metadata=metadata or {"hnsw:space": "cosine"},
            )
            self._collection_ids[name] = name
            return name
        # Ensure our main collection exists
        _ = self.collection
        return name
    
    def upsert(self, collection_name: str, ids: list, embeddings: list,
               documents: list, metadatas: list):
        """Upsert documents."""
        if collection_name != self.collection_name:
            # Get the requested collection
            collection = self.client.get_or_create_collection(name=collection_name)
        else:
            collection = self.collection
        
        collection.upsert(
            ids=ids,
            embeddings=embeddings,
            documents=documents,
            metadatas=metadatas,
        )
    
    def query(self, collection_name: str, query_embeddings: list, n_results: int = 5,
              where: dict = None, include: list = None) -> dict:
        """Query collection."""
        if collection_name != self.collection_name:
            collection = self.client.get_collection(name=collection_name)
        else:
            collection = self.collection
        
        result = collection.query(
            query_embeddings=query_embeddings,
            n_results=n_results,
            where=where,
            include=include or ["documents", "metadatas", "distances"],
        )
        return result
    
    def delete(self, collection_name: str, where: dict):
        """Delete documents."""
        if collection_name != self.collection_name:
            collection = self.client.get_collection(name=collection_name)
        else:
            collection = self.collection
        
        collection.delete(where=where)
    
    def delete_collection(self, collection_name: str):
        """Delete collection."""
        self.client.delete_collection(name=collection_name)
        if collection_name == self.collection_name:
            self._collection = None
        self._collection_ids.pop(collection_name, None)
    
    def count(self, collection_name: str) -> int:
        """Count documents."""
        if collection_name != self.collection_name:
            collection = self.client.get_collection(name=collection_name)
        else:
            collection = self.collection
        return collection.count()


class ChromaDBStore:
    """Vector store using ChromaDB - supports both HTTP (Docker) and persistent (local) clients."""

    def __init__(self, config: RAGConfig):
        self.config = config
        self._client = None
        self._collection_id = None
        self._embedding_service = EmbeddingService(config.embedding_model)
        self._use_persistent = False

    @property
    def client(self):
        """Lazy load ChromaDB client - tries HTTP first, falls back to persistent."""
        if self._client is None:
            # Try HTTP client first if configured
            http_success = False
            if self.config.chroma_use_http:
                try:
                    http_client = ChromaDBHttpClient(
                        host=self.config.chroma_host,
                        port=self.config.chroma_port,
                    )
                    # Verify connection - this will raise if ChromaDB server is not available
                    http_client.heartbeat()
                    self._client = http_client
                    self._use_persistent = False
                    http_success = True
                    logger.info(f"Connected to ChromaDB HTTP server at {self.config.chroma_host}:{self.config.chroma_port}")
                except Exception as e:
                    logger.warning(f"Failed to connect to ChromaDB HTTP server: {e}")
                    logger.info("Falling back to persistent ChromaDB client...")
            
            # Use persistent client if HTTP failed or not configured
            if not http_success:
                try:
                    self._client = ChromaDBPersistentClient(
                        persist_directory=self.config.chroma_persist_directory,
                        collection_name=self.config.chroma_collection,
                    )
                    self._use_persistent = True
                    logger.info(f"Using ChromaDB persistent client at {self.config.chroma_persist_directory}")
                except Exception as e:
                    logger.error(f"Failed to initialize ChromaDB persistent client: {e}")
                    raise
        return self._client

    @property
    def collection_name(self) -> str:
        """Get collection name and ensure it exists."""
        # Ensure client is initialized first
        _ = self.client
        
        if self._collection_id is None:
            if self._use_persistent:
                # Persistent client handles collection creation automatically
                _ = self.client.collection  # Access to trigger creation
            else:
                # HTTP client needs explicit collection creation
                self._collection_id = self.client.get_or_create_collection(
                    name=self.config.chroma_collection,
                    metadata={"hnsw:space": "cosine"},
                )
            logger.info(f"Using collection: {self.config.chroma_collection}")
        return self.config.chroma_collection

    async def add_documents(self, documents: List[Document]) -> List[str]:
        """Add documents to vector store."""
        if not documents:
            return []

        ids = [doc.id for doc in documents]
        contents = [doc.content for doc in documents]
        metadatas = [doc.metadata for doc in documents]

        # Generate embeddings
        embeddings = self._embedding_service.encode(contents).tolist()

        # Ensure collection exists
        _ = self.collection_name

        # Upsert to collection
        self.client.upsert(
            collection_name=self.config.chroma_collection,
            ids=ids,
            embeddings=embeddings,
            documents=contents,
            metadatas=metadatas,
        )

        logger.info(f"Added {len(documents)} documents to ChromaDB")
        return ids

    async def search(
        self,
        query: str,
        top_k: int = 5,
        filters: Optional[Dict[str, Any]] = None,
    ) -> List[SearchResult]:
        """Search for similar documents."""
        # Generate query embedding
        query_embedding = self._embedding_service.encode_single(query).tolist()

        # Build where clause
        where = None
        if filters:
            where = {}
            for key, value in filters.items():
                if isinstance(value, list):
                    where[key] = {"$in": value}
                else:
                    where[key] = value

        # Ensure collection exists
        _ = self.collection_name

        # Search
        results = self.client.query(
            collection_name=self.config.chroma_collection,
            query_embeddings=[query_embedding],
            n_results=top_k,
            where=where,
            include=["documents", "metadatas", "distances"],
        )

        # Convert to SearchResult
        search_results = []
        if results.get("ids") and results["ids"][0]:
            for i, doc_id in enumerate(results["ids"][0]):
                distance = results["distances"][0][i] if results.get("distances") else 0
                score = 1 - distance

                search_results.append(SearchResult(
                    id=doc_id,
                    content=results["documents"][0][i] if results.get("documents") else "",
                    score=score,
                    metadata=results["metadatas"][0][i] if results.get("metadatas") else {},
                ))

        return search_results

    async def delete_by_document_id(self, document_id: str) -> bool:
        """Delete all chunks for a document."""
        try:
            self.client.delete(
                collection_name=self.config.chroma_collection,
                where={"document_id": document_id}
            )
            return True
        except Exception as e:
            logger.error(f"Failed to delete document {document_id}: {e}")
            return False

    async def clear(self) -> bool:
        """Clear all documents."""
        try:
            self.client.delete_collection(self.config.chroma_collection)
            self._collection_id = None
            # Recreate collection
            _ = self.collection_name
            return True
        except Exception as e:
            # If collection doesn't exist (404), that's fine - it's already "cleared"
            if "404" in str(e) or "Not Found" in str(e):
                logger.info(f"Collection doesn't exist, nothing to clear")
                self._collection_id = None
                return True
            logger.error(f"Failed to clear collection: {e}")
            return False

    async def count(self) -> int:
        """Get document count."""
        try:
            return self.client.count(self.config.chroma_collection)
        except Exception:
            return 0


# ============ Document Processing ============

class DocumentProcessor:
    """Process documents into chunks."""

    SUPPORTED_EXTENSIONS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'docx',
        '.txt': 'text',
        '.md': 'markdown',
        '.html': 'html',
        '.htm': 'html',
        '.json': 'json',
        '.csv': 'csv',
    }

    def __init__(self, config: RAGConfig):
        self.config = config

    def extract_text(self, content: bytes, extension: str) -> str:
        """Extract text from document."""
        ext = extension.lower()

        if ext == '.pdf':
            return self._extract_pdf(content)
        elif ext in ['.docx', '.doc']:
            return self._extract_docx(content)
        elif ext in ['.html', '.htm']:
            return self._extract_html(content)
        else:
            return content.decode('utf-8', errors='ignore')

    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF. Falls back to OCR for scanned PDFs."""
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            text_parts = []
            pages_without_text = []

            logger.info(f"PDF has {len(reader.pages)} pages")

            # First try standard text extraction
            for i, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        text_parts.append(text.strip())
                        logger.debug(f"Page {i+1}: extracted {len(text)} chars")
                    else:
                        pages_without_text.append(i)
                except Exception as page_err:
                    logger.warning(f"Page {i+1}: extraction error - {page_err}")
                    pages_without_text.append(i)

            # If most pages have no text, try OCR
            if len(pages_without_text) > len(reader.pages) * 0.5:
                logger.info(f"PDF appears to be scanned ({len(pages_without_text)}/{len(reader.pages)} pages without text). Attempting OCR...")
                ocr_text = self._extract_pdf_ocr(content)
                if ocr_text:
                    return ocr_text
                logger.warning("OCR extraction failed or returned no text")

            result = "\n\n".join(text_parts)
            logger.info(f"Total extracted: {len(result)} chars from {len(text_parts)} pages")
            return result
        except ImportError:
            raise ImportError("pypdf not installed. Run: pip install pypdf")
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise

    def _extract_pdf_ocr(self, content: bytes) -> str:
        """Extract text from scanned PDF using OCR."""
        try:
            from pdf2image import convert_from_bytes
            import pytesseract
            from PIL import Image

            # Configure Tesseract path for Windows
            import platform
            if platform.system() == "Windows":
                tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
                if os.path.exists(tesseract_path):
                    pytesseract.pytesseract.tesseract_cmd = tesseract_path

            logger.info("Converting PDF to images for OCR...")

            # Find poppler path for Windows
            poppler_path = None
            if platform.system() == "Windows":
                # Check common installation paths
                possible_paths = [
                    r"C:\Users\JIPL\AppData\Local\Microsoft\WinGet\Packages\oschwartz10612.Poppler_Microsoft.Winget.Source_8wekyb3d8bbwe\poppler-25.07.0\Library\bin",
                    r"C:\Program Files\poppler\Library\bin",
                    r"C:\Program Files\poppler\bin",
                ]
                for path in possible_paths:
                    if os.path.exists(path):
                        poppler_path = path
                        logger.info(f"Found poppler at: {path}")
                        break

            # Convert PDF pages to images
            try:
                if poppler_path:
                    images = convert_from_bytes(content, dpi=200, poppler_path=poppler_path)
                else:
                    images = convert_from_bytes(content, dpi=200)
            except Exception as e:
                logger.warning(f"pdf2image failed (poppler may not be installed): {e}")
                logger.info("Install poppler: winget install poppler (Windows) or brew install poppler (Mac)")
                return ""

            text_parts = []
            for i, image in enumerate(images):
                try:
                    # Run OCR on each page
                    text = pytesseract.image_to_string(image)
                    if text and text.strip():
                        text_parts.append(text.strip())
                        logger.debug(f"OCR Page {i+1}: extracted {len(text)} chars")
                    else:
                        logger.debug(f"OCR Page {i+1}: no text found")
                except Exception as ocr_err:
                    logger.warning(f"OCR Page {i+1}: error - {ocr_err}")
                    continue

            result = "\n\n".join(text_parts)
            logger.info(f"OCR extracted: {len(result)} chars from {len(text_parts)} pages")
            return result

        except ImportError as e:
            logger.warning(f"OCR dependencies not installed: {e}")
            logger.info("Install: pip install pytesseract pdf2image Pillow")
            logger.info("Also install Tesseract OCR: winget install tesseract (Windows)")
            return ""
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return ""

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX."""
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))
            text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

    def _extract_html(self, content: bytes) -> str:
        """Extract text from HTML."""
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content, 'html.parser')
            for script in soup(["script", "style"]):
                script.decompose()
            return soup.get_text(separator='\n', strip=True)
        except ImportError:
            import re
            text = content.decode('utf-8', errors='ignore')
            return re.sub(r'<[^>]+>', ' ', text)

    def chunk_text(self, text: str, document_id: str, metadata: Dict[str, Any]) -> List[Document]:
        """Split text into chunks."""
        chunks = []
        paragraphs = text.split('\n\n')

        current_chunk = ""
        chunk_index = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            if len(current_chunk) + len(para) + 2 > self.config.chunk_size and current_chunk:
                # Save current chunk
                chunk_id = f"{document_id}_chunk_{chunk_index}"
                chunks.append(Document(
                    id=chunk_id,
                    content=current_chunk.strip(),
                    metadata={
                        **metadata,
                        "document_id": document_id,
                        "chunk_index": chunk_index,
                    }
                ))
                chunk_index += 1

                # Keep overlap
                if self.config.chunk_overlap > 0 and len(current_chunk) > self.config.chunk_overlap:
                    current_chunk = current_chunk[-self.config.chunk_overlap:]
                else:
                    current_chunk = ""

            current_chunk += ("\n\n" if current_chunk else "") + para

        # Add final chunk
        if current_chunk.strip():
            chunk_id = f"{document_id}_chunk_{chunk_index}"
            chunks.append(Document(
                id=chunk_id,
                content=current_chunk.strip(),
                metadata={
                    **metadata,
                    "document_id": document_id,
                    "chunk_index": chunk_index,
                }
            ))

        return chunks

    def process(
        self,
        content: bytes,
        filename: str,
        document_id: Optional[str] = None,
        extra_metadata: Optional[Dict[str, Any]] = None,
    ) -> Tuple[List[Document], ProcessingResult]:
        """Process a document into chunks."""
        start_time = time.time()

        if document_id is None:
            content_hash = hashlib.md5(content).hexdigest()[:8]
            document_id = f"doc_{content_hash}_{uuid4().hex[:8]}"

        ext = Path(filename).suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            return [], ProcessingResult(
                success=False,
                document_id=document_id,
                filename=filename,
                chunks_created=0,
                processing_time_ms=0,
                error=f"Unsupported file type: {ext}",
            )

        try:
            # Extract text
            text = self.extract_text(content, ext)

            if not text.strip():
                return [], ProcessingResult(
                    success=False,
                    document_id=document_id,
                    filename=filename,
                    chunks_created=0,
                    processing_time_ms=(time.time() - start_time) * 1000,
                    error="No text content extracted",
                )

            # Build metadata
            metadata = {
                "filename": filename,
                "file_type": self.SUPPORTED_EXTENSIONS[ext],
                "total_chars": len(text),
                **(extra_metadata or {}),
            }

            # Chunk text
            chunks = self.chunk_text(text, document_id, metadata)

            return chunks, ProcessingResult(
                success=True,
                document_id=document_id,
                filename=filename,
                chunks_created=len(chunks),
                processing_time_ms=(time.time() - start_time) * 1000,
                metadata=metadata,
            )

        except Exception as e:
            logger.exception(f"Error processing {filename}")
            return [], ProcessingResult(
                success=False,
                document_id=document_id,
                filename=filename,
                chunks_created=0,
                processing_time_ms=(time.time() - start_time) * 1000,
                error=str(e),
            )


# ============ Ray Parallel Processing ============

def init_ray_if_available(config: RAGConfig) -> Tuple[bool, None]:
    """Initialize Ray connection if available and enabled.

    Connects to Ray cluster using the ray Python package.

    Returns:
        Tuple of (is_available, None)
    """
    if not config.use_ray:
        return False, None

    try:
        import ray
        if not ray.is_initialized():
            try:
                ray.init(address=config.ray_address, ignore_reinit_error=True)
                logger.info(f"Connected to Ray cluster at {config.ray_address}")
                # Log cluster resources
                resources = ray.available_resources()
                logger.info(f"Ray cluster resources: {resources}")
                return True, None
            except Exception as connect_err:
                logger.warning(f"Failed to connect to Ray cluster: {connect_err}")
                return False, None
        else:
            logger.info("Ray already initialized")
            return True, None
    except ImportError:
        logger.warning("Ray package not installed. Run: pip install ray[client]")
        return False, None
    except Exception as e:
        logger.warning(f"Ray initialization failed: {e}")
        return False, None


def create_ray_process_task(config: RAGConfig):
    """Create Ray remote task for document processing.

    Only works if ray package is available.
    """
    try:
        import ray

        @ray.remote
        def process_document_ray(
            content: bytes,
            filename: str,
            document_id: Optional[str],
            chunk_size: int,
            chunk_overlap: int,
        ) -> Tuple[List[Dict], Dict]:
            """Ray task for processing a document."""
            processor = DocumentProcessor(RAGConfig(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
            ))
            chunks, result = processor.process(content, filename, document_id)

            # Convert to dicts for serialization
            chunk_dicts = [c.model_dump() for c in chunks]
            result_dict = {
                "success": result.success,
                "document_id": result.document_id,
                "filename": result.filename,
                "chunks_created": result.chunks_created,
                "processing_time_ms": result.processing_time_ms,
                "error": result.error,
                "metadata": result.metadata,
            }
            return chunk_dicts, result_dict

        return process_document_ray
    except ImportError:
        return None


# ============ RabbitMQ Task Queue ============

class RabbitMQQueue:
    """RabbitMQ task queue for async document processing."""

    def __init__(self, config: RAGConfig):
        self.config = config
        self._connection = None
        self._channel = None

    def connect(self):
        """Connect to RabbitMQ."""
        try:
            import pika

            credentials = pika.PlainCredentials(
                self.config.rabbitmq_user,
                self.config.rabbitmq_password,
            )
            self._connection = pika.BlockingConnection(
                pika.ConnectionParameters(
                    host=self.config.rabbitmq_host,
                    port=self.config.rabbitmq_port,
                    credentials=credentials,
                )
            )
            self._channel = self._connection.channel()
            self._channel.queue_declare(queue=self.config.rabbitmq_queue, durable=True)
            logger.info(f"Connected to RabbitMQ at {self.config.rabbitmq_host}:{self.config.rabbitmq_port}")
            return True
        except ImportError:
            logger.warning("pika not installed. Run: pip install pika")
            return False
        except Exception as e:
            logger.warning(f"Failed to connect to RabbitMQ: {e}")
            return False

    def publish(self, message: Dict[str, Any]):
        """Publish message to queue."""
        if self._channel:
            import pika
            self._channel.basic_publish(
                exchange='',
                routing_key=self.config.rabbitmq_queue,
                body=json.dumps(message),
                properties=pika.BasicProperties(delivery_mode=2),
            )

    def close(self):
        """Close connection."""
        if self._connection:
            self._connection.close()


# ============ Main RAG Pipeline ============

class IntelliBooksPipeline:
    """Main RAG Pipeline for Intellibooks Studio."""

    def __init__(self, config: Optional[RAGConfig] = None):
        self.config = config or load_config()
        self.vector_store = ChromaDBStore(self.config)
        self.processor = DocumentProcessor(self.config)
        self.embedding_service = EmbeddingService(self.config.embedding_model)

        # Initialize Ray if enabled
        self.ray_available, _ = init_ray_if_available(self.config)
        self.ray_task = create_ray_process_task(self.config) if self.ray_available else None

        # Initialize RabbitMQ if enabled
        self.rabbitmq = None
        if self.config.use_rabbitmq:
            self.rabbitmq = RabbitMQQueue(self.config)
            self.rabbitmq.connect()

    async def ingest_document(
        self,
        content: bytes,
        filename: str,
        document_id: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> ProcessingResult:
        """Ingest a single document."""
        start_time = time.time()

        # Process document
        chunks, result = self.processor.process(content, filename, document_id, metadata)

        if not result.success or not chunks:
            return result

        # Add to vector store
        await self.vector_store.add_documents(chunks)

        result.processing_time_ms = (time.time() - start_time) * 1000
        return result

    async def ingest_documents_parallel(
        self,
        documents: List[Tuple[bytes, str, Optional[str], Optional[Dict]]],
    ) -> List[ProcessingResult]:
        """Ingest multiple documents in parallel using Ray or ThreadPoolExecutor."""
        start_time = time.time()

        if self.ray_available and self.ray_task:
            return await self._ingest_with_ray(documents)
        else:
            return await self._ingest_with_threads(documents)

    async def _ingest_with_ray(
        self,
        documents: List[Tuple[bytes, str, Optional[str], Optional[Dict]]],
    ) -> List[ProcessingResult]:
        """Ingest documents using Ray for parallel processing."""
        import ray

        # Submit all tasks
        futures = []
        for content, filename, doc_id, metadata in documents:
            future = self.ray_task.remote(
                content, filename, doc_id,
                self.config.chunk_size,
                self.config.chunk_overlap,
            )
            futures.append(future)

        # Gather results
        ray_results = ray.get(futures)

        # Process results and add to vector store
        all_chunks = []
        results = []

        for chunk_dicts, result_dict in ray_results:
            chunks = [Document(**c) for c in chunk_dicts]
            all_chunks.extend(chunks)
            results.append(ProcessingResult(**result_dict))

        # Batch add to vector store
        if all_chunks:
            for i in range(0, len(all_chunks), self.config.batch_size):
                batch = all_chunks[i:i + self.config.batch_size]
                await self.vector_store.add_documents(batch)

        logger.info(f"Ray processed {len(documents)} documents ({len(all_chunks)} chunks)")
        return results

    async def _ingest_with_threads(
        self,
        documents: List[Tuple[bytes, str, Optional[str], Optional[Dict]]],
    ) -> List[ProcessingResult]:
        """Ingest documents using ThreadPoolExecutor."""
        loop = asyncio.get_event_loop()

        with ThreadPoolExecutor(max_workers=self.config.max_workers) as executor:
            # Process documents in parallel
            futures = []
            for content, filename, doc_id, metadata in documents:
                future = loop.run_in_executor(
                    executor,
                    self.processor.process,
                    content, filename, doc_id, metadata,
                )
                futures.append(future)

            # Gather results
            processed = await asyncio.gather(*futures)

        # Collect chunks and results
        all_chunks = []
        results = []

        for chunks, result in processed:
            all_chunks.extend(chunks)
            results.append(result)

        # Batch add to vector store
        if all_chunks:
            for i in range(0, len(all_chunks), self.config.batch_size):
                batch = all_chunks[i:i + self.config.batch_size]
                await self.vector_store.add_documents(batch)

        logger.info(f"Processed {len(documents)} documents ({len(all_chunks)} chunks)")
        return results

    async def query(
        self,
        query: str,
        top_k: int = 5,
        filters: Optional[Dict[str, Any]] = None,
    ) -> RAGResponse:
        """Query the knowledge base with RAG."""
        start_time = time.time()

        # Search for relevant documents
        search_results = await self.vector_store.search(query, top_k, filters)

        if not search_results:
            return RAGResponse(
                answer="I couldn't find any relevant information in the knowledge base.",
                sources=[],
                query=query,
                processing_time_ms=(time.time() - start_time) * 1000,
            )

        # Build context
        context_parts = []
        sources = []

        for i, result in enumerate(search_results):
            context_parts.append(f"[Source {i+1}]: {result.content}")
            sources.append({
                "id": result.id,
                "content_preview": result.content[:200] + "..." if len(result.content) > 200 else result.content,
                "score": result.score,
                "metadata": result.metadata,
            })

        context = "\n\n".join(context_parts)

        # Generate answer
        answer = await self._generate_answer(query, context)

        return RAGResponse(
            answer=answer,
            sources=sources,
            query=query,
            processing_time_ms=(time.time() - start_time) * 1000,
            confidence=search_results[0].score if search_results else 0,
        )

    async def _generate_answer(self, query: str, context: str) -> str:
        """Generate answer using LLM."""
        try:
            from langchain_openai import ChatOpenAI
            from langchain_core.prompts import ChatPromptTemplate
            from langchain_core.output_parsers import StrOutputParser

            if not self.config.openrouter_api_key:
                return "LLM not configured. Please set OPENROUTER_API_KEY."

            llm = ChatOpenAI(
                model=self.config.openrouter_model,
                api_key=self.config.openrouter_api_key,
                base_url="https://openrouter.ai/api/v1",
                temperature=0.3,
                default_headers={
                    "HTTP-Referer": "http://localhost:8002",
                    "X-Title": "Intellibooks Studio RAG",
                },
            )

            prompt = ChatPromptTemplate.from_messages([
                ("system", """You are Intellibooks AI, a helpful assistant for the Intellibooks Studio platform.
Answer questions based ONLY on the provided context from the knowledge base.

Guidelines:
- Be concise but comprehensive
- Cite sources when relevant (e.g., "According to Source 1...")
- If the context doesn't contain enough information, say so
- Never make up information not in the context"""),
                ("human", """Context from knowledge base:
{context}

Question: {query}

Answer:"""),
            ])

            chain = prompt | llm | StrOutputParser()
            answer = await chain.ainvoke({"context": context, "query": query})
            return answer.strip()

        except Exception as e:
            logger.exception("Error generating answer")
            return f"Error generating answer: {str(e)}"

    async def search(
        self,
        query: str,
        top_k: int = 5,
        filters: Optional[Dict[str, Any]] = None,
    ) -> List[SearchResult]:
        """Semantic search without LLM generation."""
        return await self.vector_store.search(query, top_k, filters)

    async def delete_document(self, document_id: str) -> bool:
        """Delete a document from the knowledge base."""
        return await self.vector_store.delete_by_document_id(document_id)

    async def clear(self) -> bool:
        """Clear the entire knowledge base."""
        return await self.vector_store.clear()

    async def get_stats(self) -> Dict[str, Any]:
        """Get pipeline statistics."""
        count = await self.vector_store.count()

        return {
            "total_chunks": count,
            "collection": self.config.chroma_collection,
            "embedding_model": self.config.embedding_model,
            "chroma_host": self.config.chroma_host,
            "chroma_port": self.config.chroma_port,
            "ray_enabled": self.ray_available,
            "rabbitmq_enabled": self.rabbitmq is not None,
        }
